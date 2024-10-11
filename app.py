from flask import Flask, request, render_template
from docx import Document
from werkzeug.utils import secure_filename
from collections import defaultdict
from textblob import TextBlob
import os

app = Flask(__name__)

# Allowable extensions for document upload
ALLOWED_EXTENSIONS = {'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Function to count words, highlighted words, and highlighted words by color
def count_highlighted_words_by_color(docx_file):
    doc = Document(docx_file)
    total_word_count = 0
    highlighted_word_count = 0
    color_highlighted_counts = defaultdict(int)  # A dictionary to track word counts by color

    for para in doc.paragraphs:
        for run in para.runs:
            words_in_run = run.text.split()
            total_word_count += len(words_in_run)

            if run.font.highlight_color:  # Check if the text is highlighted
                highlighted_word_count += len(words_in_run)
                highlight_color = run.font.highlight_color
                color_highlighted_counts[highlight_color] += len(words_in_run)  # Increment color-specific count

    return total_word_count, highlighted_word_count, color_highlighted_counts

# Function to analyze sentiment of the entire document
def analyze_sentiment(docx_file):
    doc = Document(docx_file)
    total_polarity = 0
    total_subjectivity = 0
    paragraph_count = 0

    for para in doc.paragraphs:
        if para.text.strip():  # Only analyze paragraphs with text
            blob = TextBlob(para.text)
            total_polarity += blob.sentiment.polarity
            total_subjectivity += blob.sentiment.subjectivity
            paragraph_count += 1

    avg_polarity = total_polarity / paragraph_count if paragraph_count > 0 else 0
    avg_subjectivity = total_subjectivity / paragraph_count if paragraph_count > 0 else 0

    return avg_polarity, avg_subjectivity

# Route for homepage
@app.route('/')
def upload_file():
    return render_template('upload.html')

# Route to handle file upload and processing
@app.route('/upload', methods=['POST'])
def upload_and_count():
    if 'file' not in request.files:
        return "No file part"
    file = request.files['file']
    if file.filename == '':
        return "No selected file"
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join('.', filename)
        file.save(filepath)
        
        # Get counts and sentiment analysis
        total_word_count, highlighted_word_count, color_highlighted_counts = count_highlighted_words_by_color(filepath)
        avg_polarity, avg_subjectivity = analyze_sentiment(filepath)
        highlighted_percentage = (highlighted_word_count / total_word_count) * 100 if total_word_count > 0 else 0
        
        # Remove the uploaded file after processing
        os.remove(filepath)
        
        # Prepare the result output
        result = f"Total Word Count: {total_word_count}\n"
        result += f"Total Highlighted Word Count: {highlighted_word_count} ({highlighted_percentage:.2f}%)\n\n"

        result += "Highlighted Word Count by Color:\n"
        for color, count in color_highlighted_counts.items():
            color_percentage = (count / total_word_count) * 100 if total_word_count > 0 else 0
            result += f"Color {color}: {count} words ({color_percentage:.2f}%)\n"
        
        result += "\nSentiment Analysis of the Document:\n"
        result += f"Average Polarity: {avg_polarity:.2f} (Polarity ranges from -1 (very negative) to 1 (very positive))\n"
        result += f"Average Subjectivity: {avg_subjectivity:.2f} (Subjectivity ranges from 0 (objective) to 1 (subjective))\n"

        return result

    return "Invalid file type. Please upload a .docx file."

if __name__ == '__main__':
    # Use the PORT environment variable provided by Railway
    port = int(os.environ.get('PORT', 8080))
    app.run(host='0.0.0.0', port=port)
